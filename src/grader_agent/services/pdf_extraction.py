"""3.2 — PDF text extraction (PyMuPDF) and DOCX plain text (python-docx)."""

from __future__ import annotations

import logging
from typing import Union

import fitz  # pymupdf
from docx import Document

from grader_agent.grading_config import pdf_max_pages
from grader_agent.models import ERROR_TYPE_VALIDATION, ErrorResult

_logger = logging.getLogger(__name__)


def _extract_text_from_docx(path: str) -> str:
    """
    Concatenate non-empty paragraph text and table cell text (row-wise).

    Merged table cells can repeat the same text once per grid cell in ``row.cells``;
    we do not deduplicate (same as typical plain-text export expectations).
    """
    doc = Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                parts.append(row_text)
    return "\n".join(parts).strip()


class PDFExtractionService:
    """Extract plain text from PDF (with page limits) or DOCX (paragraphs + tables)."""

    def extract(self, pdf_path: str, *, request_id: str | None = None) -> Union[str, ErrorResult]:
        """
        Paso 1: texto plano desde un archivo cuya ruta termina en ``.pdf`` o ``.docx``.

        ``pdf_path`` debe apuntar a un PDF válido **o** a un DOCX de Word OOXML válido según ese
        sufijo (no hay comprobación de MIME aquí).

        **PDF** — Abre con PyMuPDF, aplica ``pdf_max_pages()``, concatena ``get_text()`` por página.

        **DOCX** — Sin límite de páginas; la longitud del texto extraído no se recorta aquí ni en
        el siguiente paso de validación de contenido (regex/LLM). El tamaño del fichero suele estar
        acotado sólo por el límite de subida HTTP de la aplicación (p. ej. ``MAX_CONTENT_LENGTH``).

        Returns:
            Texto plano no vacío o ``ErrorResult`` (archivo ilegible, PDF demasiado largo por páginas,
            documento vacío o DOCX ilegible).
        """
        if request_id:
            _logger.debug("document_extract request_id=%s path=%s", request_id, pdf_path)
        low = (pdf_path or "").lower()
        if low.endswith(".docx"):
            try:
                texto = _extract_text_from_docx(pdf_path)
            except Exception as exc:
                _logger.info("DOCX open/read failed: %s", exc)
                return ErrorResult(
                    error_type=ERROR_TYPE_VALIDATION,
                    message=(
                        "El archivo no se pudo leer como Word (.docx) "
                        "(inválido, corrupto o no es DOCX). Volvé a exportar el entregable e intentá de nuevo."
                    ),
                    detail=str(exc),
                )
            if not texto:
                return ErrorResult(
                    error_type=ERROR_TYPE_VALIDATION,
                    message="El documento Word no contiene texto extraíble (vacío o sin párrafos con texto).",
                    detail=None,
                )
            return texto

        try:
            doc = fitz.open(pdf_path)
        except Exception as exc:
            _logger.info("PDF open failed: %s", exc)
            return ErrorResult(
                error_type=ERROR_TYPE_VALIDATION,
                message=(
                    "El archivo no se pudo leer como PDF (inválido, corrupto o no es PDF). "
                    "Volvé a exportar el entregable e intentá de nuevo."
                ),
                detail=str(exc),
            )

        try:
            n = len(doc)
            max_pages = pdf_max_pages()
            if n > max_pages:
                return ErrorResult(
                    error_type=ERROR_TYPE_VALIDATION,
                    message=f"El PDF tiene {n} páginas. El máximo permitido es {max_pages}.",
                    detail=None,
                )
            parts: list[str] = []
            for pagina in doc:
                parts.append(pagina.get_text())
            texto = "\n".join(parts).strip()
            if not texto:
                return ErrorResult(
                    error_type=ERROR_TYPE_VALIDATION,
                    message="El PDF no contiene texto extraíble (vacío o solo imágenes).",
                    detail=None,
                )
            return texto
        finally:
            doc.close()


__all__ = ["PDFExtractionService"]
